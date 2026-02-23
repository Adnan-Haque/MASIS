import io
import json
import logging
from PyPDF2 import PdfReader
from docx import Document
from app.ingestion.vision_processor import process_image_with_vision

logger = logging.getLogger(__name__)

MAX_CHARS = 1500      # safe embedding size
OVERLAP_CHARS = 200   # chars carried from end of previous chunk into next


def split_large_text(text, chunk_type="text"):
    """
    Split text into overlapping chunks of MAX_CHARS.

    CHANGE vs original:
    - After yielding a chunk, the last OVERLAP_CHARS of that chunk
      are prepended to the next chunk's buffer.
    - This ensures sentences/figures that straddle a chunk boundary
      appear in BOTH adjacent chunks, so retrieval can find them
      regardless of which chunk is returned.
    - Everything else (paragraph splitting logic, yield structure,
      chunk dict shape) is identical to the original.
    """
    paragraphs = text.split("\n")
    buffer = ""

    for para in paragraphs:
        if len(buffer) + len(para) < MAX_CHARS:
            buffer += para + "\n"
        else:
            if buffer.strip():
                yield {
                    "chunk_type": chunk_type,
                    "text": buffer.strip(),
                    "structured_data": None,
                }

            # CHANGE: carry last OVERLAP_CHARS into next buffer
            # Original was: buffer = para + "\n"  (hard reset, zero overlap)
            overlap = buffer[-OVERLAP_CHARS:] if len(buffer) > OVERLAP_CHARS else buffer
            buffer = overlap + para + "\n"

    if buffer.strip():
        yield {
            "chunk_type": chunk_type,
            "text": buffer.strip(),
            "structured_data": None,
        }


def extract_text_stream(filename, file_bytes):

    filename_lower = filename.lower()

    # ================= IMAGE =================
    if filename_lower.endswith((".png", ".jpg", ".jpeg")):
        try:
            summary, structured = process_image_with_vision(file_bytes)

            yield {
                "chunk_type": structured.get("type", "image_text"),
                "text": summary,
                "structured_data": structured,
            }
        except Exception as e:
            logger.error(f"Vision failed for image {filename}: {str(e)}")
        return

    # ================= PDF =================
    if filename_lower.endswith(".pdf"):
        try:
            reader = PdfReader(io.BytesIO(file_bytes))

            for page_index, page in enumerate(reader.pages):
                text = page.extract_text()

                if text and text.strip():
                    for chunk in split_large_text(text, "text"):
                        chunk["page_number"] = page_index
                        yield chunk

        except Exception as e:
            logger.error(f"PDF extraction failed: {str(e)}")
        return

    # ================= DOCX =================
    if filename_lower.endswith(".docx"):
        try:
            doc = Document(io.BytesIO(file_bytes))

            # CHANGE: group paragraphs in sliding windows of 3 with 1 overlap
            # instead of yielding one paragraph at a time.
            #
            # Original:
            #   for para in doc.paragraphs:
            #       if para.text.strip():
            #           yield { "text": para.text, ... }
            #
            # Problem: a single paragraph like "Reduce churn by 15%." has no
            # surrounding context so it scores poorly in retrieval.
            #
            # Fix: combine 3 paragraphs per chunk, slide by 2 (1 overlap),
            # so each chunk shares one paragraph with its neighbours.
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

            WINDOW = 3
            STEP = 2  # WINDOW - 1 overlap paragraph

            i = 0
            while i < len(paragraphs):
                window = paragraphs[i: i + WINDOW]
                combined = "\n".join(window)

                if len(combined) > MAX_CHARS:
                    # window too large — fall back to character splitter
                    for chunk in split_large_text(combined, "text"):
                        yield chunk
                else:
                    yield {
                        "chunk_type": "text",
                        "text": combined,
                        "structured_data": None,
                    }

                i += STEP

            # Tables — unchanged from original
            for table_index, table in enumerate(doc.tables):
                for row_index, row in enumerate(table.rows):
                    row_text = " | ".join(cell.text for cell in row.cells)

                    if row_text.strip():
                        yield {
                            "chunk_type": "table_row",
                            "text": row_text,
                            "structured_data": None,
                            "table_index": table_index,
                            "row_index": row_index,
                        }

        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
        return

    # ================= JSON =================
    if filename_lower.endswith(".json"):
        try:
            text = file_bytes.decode("utf-8")
            data = json.loads(text)

            if isinstance(data, list):
                for idx, item in enumerate(data):
                    yield {
                        "chunk_type": "json_item",
                        "text": json.dumps(item),
                        "structured_data": item,
                        "item_index": idx,
                    }

            elif isinstance(data, dict):
                for key, value in data.items():
                    yield {
                        "chunk_type": "json_field",
                        "text": f"{key}: {json.dumps(value)}",
                        "structured_data": {key: value},
                    }

        except Exception as e:
            logger.error(f"JSON decode failed: {str(e)}")
        return

    # ================= XML =================
    if filename_lower.endswith(".xml"):
        try:
            text = file_bytes.decode("utf-8")
            for chunk in split_large_text(text, "xml_block"):
                yield chunk
        except Exception as e:
            logger.error(f"XML decode failed: {str(e)}")
        return

    # ================= TXT / FALLBACK =================
    try:
        text = file_bytes.decode("utf-8", errors="ignore")

        if text.strip():
            for chunk in split_large_text(text):
                yield chunk

    except Exception as e:
        logger.error(f"Fallback decode failed: {str(e)}")